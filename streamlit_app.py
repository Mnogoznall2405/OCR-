import streamlit as st
import pandas as pd
import numpy as np
import io
import os
import json
import time
import base64
import requests
from PIL import Image
from pathlib import Path
from dotenv import load_dotenv
from fpdf import FPDF
from docx import Document
import matplotlib.pyplot as plt
from wordcloud import WordCloud
import shutil
import datetime

# Загружаем переменные из .env файла если он существует
load_dotenv()

# Константы
CACHE_DIR = Path('cache')
HISTORY_DIR = Path('history')
STATS_FILE = Path('stats.json')
INPUT_WATCH_DIR = Path('input_watch')
OUTPUT_RESULTS_DIR = Path('output_results')

# Создаем необходимые директории
CACHE_DIR.mkdir(exist_ok=True)
HISTORY_DIR.mkdir(exist_ok=True)
INPUT_WATCH_DIR.mkdir(exist_ok=True)
OUTPUT_RESULTS_DIR.mkdir(exist_ok=True)

# Настройка страницы
st.set_page_config(
    page_title="OCR и Перевод текста",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Получение API ключа из переменных окружения или Streamlit secrets
def get_api_key():
    # Сначала проверяем переменные окружения
    api_key = os.getenv("OCR_API_KEY")
    
    # Если не найдено в переменных окружения, проверяем Streamlit secrets
    if not api_key and hasattr(st, "secrets"):
        api_key = st.secrets.get("OCR_API_KEY")
    
    return api_key

# Глобальные переменные
if 'OCR_API_KEY' not in st.session_state:
    # Получаем ключ из защищенного источника
    api_key = get_api_key()
    if api_key:
        st.session_state.OCR_API_KEY = api_key
    else:
        # Показываем форму для ввода API ключа, если он не найден
        st.session_state.OCR_API_KEY = None

OCR_SPACE_URL = 'https://api.ocr.space/parse/image'

# Словарь поддерживаемых языков
SUPPORTED_LANGUAGES = {
    'en': 'Английский',
    'ru': 'Русский',
    'de': 'Немецкий',
    'fr': 'Французский',
    'es': 'Испанский',
    'it': 'Итальянский',
    'pt': 'Португальский',
    'nl': 'Нидерландский',
    'pl': 'Польский',
    'uk': 'Украинский',
    'ja': 'Японский',
    'ko': 'Корейский',
    'zh': 'Китайский',
    'ar': 'Арабский'
}

# Соответствие языков для OCR.space
OCR_SPACE_LANGUAGES = {
    'en': 'eng',
    'ru': 'rus',
    'de': 'ger',
    'fr': 'fre',
    'es': 'spa',
    'it': 'ita',
    'pt': 'por',
    'nl': 'dut',
    'pl': 'pol',
    'uk': 'ukr',
    'ja': 'jpn',
    'ko': 'kor',
    'zh': 'chi_sim',
    'ar': 'ara'
}

# Функция перевода текста
def translate_text(text, source_lang='auto', target_lang='en'):
    try:
        # Автоматически определяем направление перевода
        if any(char in 'абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ' for char in text):
            source_lang = 'ru'
            target_lang = 'en'
        else:
            source_lang = 'en'
            target_lang = 'ru'
        
        # Используем бесплатный API для перевода
        url = "https://translate.googleapis.com/translate_a/single"
        params = {
            "client": "gtx",
            "sl": source_lang,
            "tl": target_lang,
            "dt": "t",
            "q": text
        }
        
        response = requests.get(url, params=params)
        if response.status_code == 200:
            result = response.json()
            translated_text = ''.join([sentence[0] for sentence in result[0]])
            return translated_text
        else:
            return f"Ошибка перевода: {response.status_code}"
    except Exception as e:
        return f"Ошибка перевода: {str(e)}"

# Функция распознавания текста с помощью OCR.space API
def ocr_space_recognize(image_data, language='auto', enhance_contrast=False, remove_noise=False):
    try:
        # Проверяем, что API ключ существует
        if not st.session_state.OCR_API_KEY:
            raise Exception("API ключ OCR.space не настроен. Пожалуйста, настройте его в настройках.")
            
        # Определяем тип файла
        file_type = 'png'  # По умолчанию PNG
        if image_data.startswith(b'\xFF\xD8\xFF'):
            file_type = 'jpg'
        elif image_data.startswith(b'%PDF'):
            file_type = 'pdf'
        
        # Кодируем изображение в base64 с правильным префиксом
        base64_image = f"data:image/{file_type};base64,{base64.b64encode(image_data).decode('utf-8')}"
        
        # Автоматическое определение языка (OCR.space будет определять автоматически)
        ocr_space_lang = 'eng' if language == 'en' else 'auto'
        
        payload = {
            'base64Image': base64_image,
            'language': ocr_space_lang,
            'isOverlayRequired': False,
            'OCREngine': 2,  # 2 - лучший движок
            'filetype': file_type.upper(),
            'detectOrientation': True,
            'scale': True,
            'isCreateSearchablePdf': False,
            'isSearchablePdfHideTextLayer': False,
            'isTable': False
        }
        
        headers = {
            'apikey': st.session_state.OCR_API_KEY,
            'Content-Type': 'application/x-www-form-urlencoded'
        }
        
        response = requests.post(OCR_SPACE_URL, data=payload, headers=headers, timeout=60)
        
        if response.status_code == 401 or response.status_code == 403:
            # Неверный API ключ
            st.session_state.OCR_API_KEY = None  # Сбрасываем ключ
            raise Exception("Неверный API ключ OCR.space. Пожалуйста, проверьте ваш ключ.")
            
        elif response.status_code != 200:
            raise Exception(f"Ошибка API OCR.space: {response.text}")
            
        result = response.json()
        if result.get('IsErroredOnProcessing'):
            error_msg = result.get('ErrorMessage', 'Неизвестная ошибка OCR')
            if "Unauthorized request" in error_msg:
                st.session_state.OCR_API_KEY = None  # Сбрасываем ключ
                raise Exception("Неверный или просроченный API ключ OCR.space")
            raise Exception(error_msg)
            
        text = result['ParsedResults'][0]['ParsedText']
        
        # Пытаемся получить определенный язык, но OCR.space его не возвращает напрямую
        # Поэтому определяем на основе содержимого текста
        detected_language = None
        
        return text.strip(), detected_language
        
    except Exception as e:
        raise Exception(f"Ошибка при распознавании: {str(e)}")

# Функции для работы с кэшем
def get_cache_key(image_data: bytes) -> str:
    import hashlib
    return hashlib.md5(image_data).hexdigest()

def save_to_cache(image_data: bytes, result: dict):
    cache_key = get_cache_key(image_data)
    cache_file = CACHE_DIR / f"{cache_key}.json"
    with open(cache_file, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

def get_from_cache(image_data: bytes) -> dict:
    cache_key = get_cache_key(image_data)
    cache_file = CACHE_DIR / f"{cache_key}.json"
    if cache_file.exists():
        with open(cache_file, 'r', encoding='utf-8') as f:
            return json.load(f)
    return None

# Функции для работы со статистикой
def update_stats(success: bool, file_size: int):
    if not STATS_FILE.exists():
        stats = {
            'total_processed': 0,
            'total_success': 0,
            'total_failed': 0,
            'total_size': 0,
            'last_processed': None
        }
    else:
        with open(STATS_FILE, 'r', encoding='utf-8') as f:
            stats = json.load(f)
    
    stats['total_processed'] += 1
    if success:
        stats['total_success'] += 1
    else:
        stats['total_failed'] += 1
    stats['total_size'] += file_size
    stats['last_processed'] = time.strftime("%Y-%m-%d %H:%M:%S")
    
    # Сохраняем статистику
    with open(STATS_FILE, 'w', encoding='utf-8') as f:
        json.dump(stats, f, ensure_ascii=False, indent=2)
    
    return stats

# Функции для проверки безопасности
def is_allowed_file(filename: str) -> bool:
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'png', 'jpg', 'jpeg', 'pdf'}

def check_file_size(file_data: bytes) -> bool:
    # Изменяем ограничение с 10 МБ до 1 МБ
    return len(file_data) <= 1 * 1024 * 1024  # 1MB

def optimize_image(image_data: bytes) -> bytes:
    try:
        # Проверяем, является ли файл PDF
        if image_data.startswith(b'%PDF'):
            # Для PDF просто возвращаем исходные данные
            return image_data
            
        img = Image.open(io.BytesIO(image_data))
        
        # Сжимаем изображение, если оно слишком большое
        if img.size[0] > 2000 or img.size[1] > 2000:
            img.thumbnail((2000, 2000), Image.Resampling.LANCZOS)
        
        # Конвертируем RGBA в RGB, если необходимо
        if img.mode == 'RGBA':
            # Создаем новое RGB изображение с белым фоном
            rgb_img = Image.new('RGB', img.size, (255, 255, 255))
            # Композиция с прозрачностью поверх белого фона
            rgb_img.paste(img, mask=img.split()[3])  # Используем альфа-канал как маску
            img = rgb_img
        elif img.mode != 'RGB':
            # Конвертируем другие режимы (например, 'P' или 'L') в RGB
            img = img.convert('RGB')
            
        # Конвертируем в JPEG для уменьшения размера
        output = io.BytesIO()
        img.save(output, format='JPEG', quality=85)
        return output.getvalue()
    except Exception as e:
        st.error(f"Ошибка при оптимизации изображения: {str(e)}")
        # Возвращаем исходные данные в случае ошибки
        return image_data

# Функция для обработки изображения
def process_image(image_data: bytes, settings: dict) -> dict:
    try:
        # Проверяем кэш, если включен
        if settings.get('use_cache', True):
            cached_result = get_from_cache(image_data)
            if cached_result:
                return cached_result
        
        # Оптимизируем изображение, если включено
        if settings.get('optimize', True):
            image_data = optimize_image(image_data)
        
        # Распознаем текст напрямую через API
        text, _ = ocr_space_recognize(
            image_data, 
            language='auto',  # Автоматическое определение языка
            enhance_contrast=settings.get('enhance_contrast', False),
            remove_noise=settings.get('remove_noise', False)
        )
        
        # Определяем язык на основе содержимого текста
        detected_language = 'en'  # По умолчанию английский
        
        # Если в тексте есть кириллические символы, считаем что язык русский
        if any(char in 'абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ' for char in text):
            detected_language = 'ru'
            
        # Подсчитываем кириллические и латинские символы для более точного определения
        cyrillic_chars = sum(1 for char in text if char in 'абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ')
        latin_chars = sum(1 for char in text if 'a' <= char.lower() <= 'z')
        
        # Если кириллицы больше, чем латиницы - русский язык
        if cyrillic_chars > latin_chars:
            detected_language = 'ru'
        
        result = {
            'text': text,
            'processing_time': f'{time.time() - settings.get("start_time", time.time()):.2f} секунд',
            'language': SUPPORTED_LANGUAGES.get(detected_language, 'Автоопределение'),
            'detected_language': detected_language
        }
            
        # Сохраняем в кэш, если включен
        if settings.get('use_cache', True):
            save_to_cache(image_data, result)
                
        return result
    except Exception as e:
        return {'error': f'Ошибка при обработке: {str(e)}'}

# Функция для сохранения в историю
def save_to_history(image_data: bytes, text: str, language: str, processing_time: str, translated_text=None, target_lang=None):
    timestamp = time.strftime("%Y%m%d-%H%M%S")
    history_file = HISTORY_DIR / f"{timestamp}.json"
    
    # Сохраняем изображение
    image_file = HISTORY_DIR / f"{timestamp}.jpg"
    try:
        if not image_data.startswith(b'%PDF'):
            img = Image.open(io.BytesIO(image_data))
            
            # Конвертируем RGBA в RGB, если необходимо
            if img.mode == 'RGBA':
                # Создаем новое RGB изображение с белым фоном
                rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                # Композиция с прозрачностью поверх белого фона
                rgb_img.paste(img, mask=img.split()[3])  # Используем альфа-канал как маску
                img = rgb_img
            elif img.mode != 'RGB':
                # Конвертируем другие режимы (например, 'P' или 'L') в RGB
                img = img.convert('RGB')
                
            img.save(image_file, format='JPEG', quality=85)
        else:
            # Для PDF сохраняем как есть
            with open(HISTORY_DIR / f"{timestamp}.pdf", 'wb') as f:
                f.write(image_data)
    except Exception as e:
        st.error(f"Ошибка при сохранении изображения: {str(e)}")
    
    # Сохраняем метаданные
    history_data = {
        'timestamp': timestamp,
        'text': text,
        'language': language,
        'processing_time': processing_time
    }
    
    # Добавляем информацию о переводе, если он был выполнен
    if translated_text and target_lang:
        history_data['translated_text'] = translated_text
        history_data['target_language'] = target_lang
    
    with open(history_file, 'w', encoding='utf-8') as f:
        json.dump(history_data, f, ensure_ascii=False, indent=2)

# Функция для загрузки истории
@st.cache_data
def load_history():
    history = []
    for history_file in sorted(HISTORY_DIR.glob('*.json'), reverse=True):
        try:
            with open(history_file, 'r', encoding='utf-8') as f:
                history_data = json.load(f)
                
                # Проверяем, существует ли соответствующее изображение
                timestamp = history_data.get('timestamp', '')
                image_file = HISTORY_DIR / f"{timestamp}.jpg"
                pdf_file = HISTORY_DIR / f"{timestamp}.pdf"
                
                if image_file.exists():
                    history_data['image_path'] = str(image_file)
                elif pdf_file.exists():
                    history_data['pdf_path'] = str(pdf_file)
                
                history.append(history_data)
        except Exception as e:
            st.error(f"Ошибка при загрузке истории: {str(e)}")
    return history

# Функция для загрузки статистики
@st.cache_data
def load_stats():
    if STATS_FILE.exists():
        with open(STATS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {
        'total_processed': 0,
        'total_success': 0,
        'total_failed': 0,
        'total_size': 0,
        'last_processed': None
    }

# Словарь языков для перевода
TRANSLATION_LANGUAGES = {
    'en': 'Английский',
    'ru': 'Русский',
    'de': 'Немецкий',
    'fr': 'Французский',
    'es': 'Испанский',
    'it': 'Итальянский',
    'zh': 'Китайский',
    'ja': 'Японский',
    'ko': 'Корейский',
    'ar': 'Арабский'
}

# Функция для экспорта текста в TXT
def export_to_txt(text: str) -> str:
    timestamp = time.strftime("%Y%m%d-%H%M%S")
    export_file = f"export_{timestamp}.txt"
    
    with open(export_file, 'w', encoding='utf-8') as f:
        f.write(text)
    
    return export_file

# Функция для экспорта текста в PDF
def export_to_pdf(text: str, translated_text: str = None) -> str:
    timestamp = time.strftime("%Y%m%d-%H%M%S")
    export_file = f"export_{timestamp}.pdf"
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, "Оригинальный текст:\n" + text)
    if translated_text:
        pdf.ln(5)
        pdf.multi_cell(0, 10, "Перевод:\n" + translated_text)
    pdf.output(export_file)
    return export_file

# Функция для экспорта текста в DOCX
def export_to_docx(text: str, translated_text: str = None) -> str:
    timestamp = time.strftime("%Y%m%d-%H%M%S")
    export_file = f"export_{timestamp}.docx"
    doc = Document()
    doc.add_heading('Оригинальный текст', level=2)
    doc.add_paragraph(text)
    if translated_text:
        doc.add_heading('Перевод', level=2)
        doc.add_paragraph(translated_text)
    doc.save(export_file)
    return export_file

# Функция для применения CSS-стилей
def apply_custom_css():
    st.markdown("""
    <style>
    /* Основные цвета и стили */
    :root {
        --primary-color: #4361ee;
        --secondary-color: #3f37c9;
        --accent-color: #4cc9f0;
        --text-color: #333;
        --light-bg: #f8f9fa;
        --dark-bg: #212529;
        --success-color: #4ade80;
        --warning-color: #fbbf24;
        --error-color: #f87171;
    }
    
    /* Стили для заголовков */
    h1, h2, h3 {
        color: var(--primary-color);
        font-weight: 600;
    }
    
    h1 {
        border-bottom: 2px solid var(--primary-color);
        padding-bottom: 0.5rem;
        margin-bottom: 1.5rem;
    }
    
    /* Стили для контейнеров */
    .result-container {
        background-color: var(--light-bg);
        padding: 1.2rem;
        border-radius: 8px;
        border-left: 4px solid var(--primary-color);
        margin-top: 1.5rem;
        box-shadow: 0 2px 5px rgba(0,0,0,0.05);
    }
    
    .translated-container {
        background-color: #e6f3ff;
        padding: 1.2rem;
        border-radius: 8px;
        border-left: 4px solid var(--accent-color);
        margin-top: 1rem;
        box-shadow: 0 2px 5px rgba(0,0,0,0.05);
    }
    
    /* Улучшение отзывчивости textarea */
    .stTextArea textarea {
        border-radius: 6px;
        border: 1px solid #ccc;
        font-size: 1rem;
    }
    
    /* Стили для предпросмотра изображений */
    .preview-container {
        padding: 1rem;
        border: 1px dashed #ccc;
        border-radius: 8px;
        text-align: center;
    }
    
    .preview-image {
        max-height: 300px;
        width: auto;
        margin: 0 auto;
        border-radius: 6px;
    }
    
    /* Улучшение стилей боковой панели */
    .css-1l4w6pd {
        background-color: #f8f9fa;
    }
    
    /* Улучшенные стили кнопок */
    .stButton button {
        border-radius: 6px;
        font-weight: 500;
        transition: all 0.2s ease;
    }
    
    .stButton button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
    }
    
    /* Стили для уведомлений */
    div[data-testid="stNotification"] {
        border-radius: 8px;
        box-shadow: 0 4px 10px rgba(0,0,0,0.1);
    }
    
    /* Стили для индикатора загрузки */
    .stSpinner svg {
        animation: spin 1s linear infinite;
    }
    
    /* Отменяем анимацию для текста в сообщении загрузки */
    .stSpinner p, .stSpinner span, .stSpinner div:not([role="progressbar"]) {
        animation: none !important;
        transform: none !important;
    }
    
    @keyframes spin {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }
    
    /* Стили для вкладок */
    .stTabs [data-baseweb="tab-list"] {
        gap: 2rem;
    }
    
    .stTabs [data-baseweb="tab"] {
        border-radius: 4px 4px 0 0;
        padding: 0.5rem 1rem;
        font-weight: 500;
    }
    
    .stTabs [aria-selected="true"] {
        background-color: var(--primary-color);
        color: white;
    }
    
    /* Улучшение внешнего вида карточек */
    .card {
        padding: 1.5rem;
        border-radius: 8px;
        box-shadow: 0 4px 8px rgba(0,0,0,0.08);
        margin-bottom: 1rem;
        transition: all 0.3s ease;
    }
    
    .card:hover {
        transform: translateY(-4px);
        box-shadow: 0 6px 12px rgba(0,0,0,0.1);
    }
    
    @media (max-width: 900px) {
        .stTabs [data-baseweb="tab-list"] {
            gap: 0.5rem;
        }
        .stTabs [data-baseweb="tab"] {
            font-size: 1rem;
            padding: 0.3rem 0.7rem;
        }
        .stButton button, .stDownloadButton button {
            font-size: 1.1rem;
            padding: 0.7rem 1.2rem;
        }
        .stTextArea textarea {
            font-size: 1.1rem;
        }
        .result-container, .translated-container, .card {
            padding: 0.7rem;
        }
        .preview-image {
            max-height: 180px;
        }
    }
    @media (max-width: 600px) {
        .stTabs [data-baseweb="tab-list"] {
            flex-direction: column;
            gap: 0.2rem;
        }
        .stTabs [data-baseweb="tab"] {
            width: 100%;
            font-size: 1rem;
        }
        .stButton button, .stDownloadButton button {
            width: 100%;
            font-size: 1.2rem;
            padding: 1rem 0.5rem;
        }
        .stTextArea textarea {
            font-size: 1.2rem;
            min-height: 120px;
        }
        .result-container, .translated-container, .card {
            padding: 0.5rem;
            margin-top: 0.7rem;
        }
        .preview-image {
            max-height: 120px;
        }
        .stColumns {
            flex-direction: column !important;
        }
    }
    </style>
    """, unsafe_allow_html=True)

# Функция для анализа текста и получения статистики
def analyze_text(text: str) -> dict:
    if not text:
        return {
            "chars_count": 0,
            "words_count": 0,
            "lines_count": 0,
            "paragraphs_count": 0,
            "letters_count": 0,
            "digits_count": 0,
            "spaces_count": 0,
            "punctuation_count": 0,
            "common_words": []
        }
    
    import re
    import string
    from collections import Counter
    
    # Общее количество символов
    chars_count = len(text)
    
    # Количество строк
    lines = text.split('\n')
    lines_count = len(lines)
    
    # Количество абзацев (непустые строки после пустой строки)
    paragraphs = [p for p in re.split(r'\n\s*\n', text) if p.strip()]
    paragraphs_count = len(paragraphs)
    
    # Количество слов
    words = re.findall(r'\b\w+\b', text.lower())
    words_count = len(words)
    
    # Подсчет букв, цифр, пробелов и знаков пунктуации
    letters_count = sum(c.isalpha() for c in text)
    digits_count = sum(c.isdigit() for c in text)
    spaces_count = sum(c.isspace() for c in text)
    punctuation_count = sum(c in string.punctuation for c in text)
    
    # Наиболее распространенные слова (исключая стоп-слова)
    stop_words = set(['и', 'в', 'на', 'с', 'по', 'для', 'не', 'от', 'за', 'к', 'а', 'the', 'and', 'of', 'to', 'in', 'a', 'is', 'that', 'for', 'on', 'with', 'by', 'at', 'as'])
    filtered_words = [word for word in words if word not in stop_words and len(word) > 2]
    word_counts = Counter(filtered_words)
    common_words = word_counts.most_common(5)  # 5 самых распространенных слов
    
    return {
        "chars_count": chars_count,
        "words_count": words_count,
        "lines_count": lines_count,
        "paragraphs_count": paragraphs_count,
        "letters_count": letters_count,
        "digits_count": digits_count,
        "spaces_count": spaces_count,
        "punctuation_count": punctuation_count,
        "common_words": common_words
    }

def show_text_visualizations(text_stats, text, key_prefix=""):
    # Облако слов
    words = [w for w in text.lower().split() if len(w) > 2]
    if words:
        wordcloud = WordCloud(width=400, height=200, background_color='white').generate(' '.join(words))
        st.markdown("#### ☁️ Облако слов")
        fig_wc, ax_wc = plt.subplots(figsize=(6, 3))
        ax_wc.imshow(wordcloud, interpolation='bilinear')
        ax_wc.axis('off')
        st.pyplot(fig_wc, clear_figure=True)
    # Круговая диаграмма
    st.markdown("#### 🥧 Распределение символов")
    labels = ['Буквы', 'Цифры', 'Пробелы', 'Пунктуация']
    sizes = [text_stats['letters_count'], text_stats['digits_count'], text_stats['spaces_count'], text_stats['punctuation_count']]
    fig_pie, ax_pie = plt.subplots()
    ax_pie.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
    ax_pie.axis('equal')
    st.pyplot(fig_pie, clear_figure=True)
    # Гистограмма длины слов
    st.markdown("#### 📊 Длина слов")
    import re
    word_lengths = [len(w) for w in re.findall(r'\b\w+\b', text)]
    if word_lengths:
        fig_hist, ax_hist = plt.subplots()
        ax_hist.hist(word_lengths, bins=range(1, max(word_lengths)+2), color='#4361ee', edgecolor='black')
        ax_hist.set_xlabel('Длина слова')
        ax_hist.set_ylabel('Количество')
        st.pyplot(fig_hist, clear_figure=True)

# Основной интерфейс приложения
def main():
    # Применяем CSS-стили
    apply_custom_css()
    
    # Проверяем наличие API ключа
    if st.session_state.OCR_API_KEY is None:
        st.title("📝 OCR Распознавание и Перевод текста")
        
        st.error("⚠️ API ключ для OCR.space не найден. Пожалуйста, настройте его для работы приложения.")
        
        with st.form("api_key_form"):
            st.write("### Введите API ключ OCR.space")
            api_key = st.text_input("API ключ", type="password", 
                                    placeholder="Введите ваш ключ API...",
                                    help="Ключ можно получить бесплатно на сайте OCR.space")
            submitted = st.form_submit_button("Сохранить", use_container_width=True)
            
            if submitted and api_key:
                st.session_state.OCR_API_KEY = api_key
                st.success("✅ API ключ успешно сохранен в текущей сессии!")
                time.sleep(1)  # Небольшая задержка для отображения сообщения
                st.rerun()
        
        st.markdown("""
        <div class="card">
            <h3>📋 Как настроить API ключ</h3>
            <ol>
                <li><b>Через .env файл</b>:
                    <ul>
                        <li>Создайте файл <code>.env</code> в корневой директории проекта</li>
                        <li>Добавьте строку: <code>OCR_API_KEY=ваш_ключ_api</code></li>
                    </ul>
                </li>
                <br>
                <li><b>Через Streamlit Secrets</b> (рекомендуется для деплоя):
                    <ul>
                        <li>Создайте файл <code>.streamlit/secrets.toml</code></li>
                        <li>Добавьте строку: <code>OCR_API_KEY="ваш_ключ_api"</code></li>
                    </ul>
                </li>
            </ol>
            <p>Получить API ключ можно на сайте <a href="https://ocr.space/ocrapi" target="_blank">OCR.space</a> (бесплатно)</p>
        </div>
        """, unsafe_allow_html=True)
        
        return  # Прерываем выполнение, пока не будет ключа
    
    # Заголовок приложения
    st.title("📝 OCR Распознавание и Перевод текста")
    st.markdown("Загрузите изображение или PDF файл для извлечения и перевода текста.")
    
    # Боковая панель
    with st.sidebar:
        st.header("⚙️ Настройки")
        
        # Автоматическое определение языка
        st.subheader("🔤 Распознавание")
        st.info("Язык текста определяется автоматически")
        
        # Расширенные настройки
        with st.expander("🛠️ Расширенные настройки"):
            optimize = st.checkbox("✅ Оптимизировать изображение", value=True, 
                                  help="Автоматически оптимизирует размер и качество изображения")
            use_cache = st.checkbox("📦 Использовать кэш", value=True,
                                   help="Сохраняет результаты распознавания для повторного использования")
            
            # Дополнительные опции
            enhance_contrast = st.checkbox("🔎 Улучшить контрастность", value=False,
                                         help="Может помочь при распознавании низкоконтрастных изображений")
            remove_noise = st.checkbox("🧹 Удалить шум", value=False,
                                     help="Может улучшить распознавание зашумленных изображений")
        
        # Информация о статистике
        st.divider()
        st.subheader("📊 Статистика")
        stats = load_stats()
        
        col1, col2 = st.columns(2)
        col1.metric("Всего обработано", stats['total_processed'])
        col2.metric("Успешно", stats['total_success'])
        
        col1, col2 = st.columns(2)
        col1.metric("Ошибок", stats['total_failed'])
        col2.metric("Размер (МБ)", round(stats['total_size'] / (1024 * 1024), 2))
        
        if stats['last_processed']:
            st.caption(f"Последняя обработка: {stats['last_processed']}")
    
    # Основная область
    tab1, tab2 = st.tabs(["📸 Распознавание", "📜 История"])
    
    # Вкладка распознавания
    with tab1:
        # Кнопка для пакетной обработки из папки
        if st.button('📂 Обработать все файлы из папки input_watch', use_container_width=True):
            files = list(INPUT_WATCH_DIR.glob('*'))
            if not files:
                st.info('В папке input_watch нет файлов для обработки.')
            else:
                st.info(f'Найдено файлов: {len(files)}. Начинаем обработку...')
                for file in files:
                    try:
                        with open(file, 'rb') as f:
                            file_data = f.read()
                        if not check_file_size(file_data):
                            st.warning(f'{file.name}: файл слишком большой, пропущен.')
                            continue
                        start_time = time.time()
                        settings = {
                            'language': 'en',
                            'optimize': True,
                            'use_cache': True,
                            'enhance_contrast': False,
                            'remove_noise': False,
                            'start_time': start_time
                        }
                        result = process_image(file_data, settings)
                        processing_time = f"{time.time() - start_time:.2f} сек."
                        success = 'error' not in result
                        update_stats(success, len(file_data))
                        if success:
                            original_text = result['text']
                            detected_language = result.get('detected_language', 'auto')
                            # Сохраняем результат в output_results
                            out_name = OUTPUT_RESULTS_DIR / f"{file.stem}_result.txt"
                            with open(out_name, 'w', encoding='utf-8') as out_f:
                                out_f.write(original_text)
                            st.success(f'{file.name}: успешно обработан!')
                        else:
                            st.error(f'{file.name}: ошибка обработки.')
                        # Перемещаем файл в архив
                        shutil.move(str(file), str(INPUT_WATCH_DIR / f"_done_{file.name}"))
                    except Exception as e:
                        st.error(f'{file.name}: {str(e)}')
        
        # Колонки для разделения загрузки и предпросмотра
        col_upload, col_preview = st.columns([2, 1])
        
        with col_upload:
            # Загрузка файлов (множественный выбор)
            uploaded_files = st.file_uploader("Выберите изображения или PDF файлы", 
                                        type=['png', 'jpg', 'jpeg', 'pdf'],
                                        help="Поддерживаются форматы PNG, JPG и PDF",
                                        label_visibility="collapsed",
                                        accept_multiple_files=True)
            
            st.caption("Лимит 1МБ на файл • PNG, JPG, JPEG, PDF")
            
            if uploaded_files:
                # Кнопка для запуска пакетной обработки
                if st.button("🔍 Распознать все файлы", type="primary", use_container_width=True):
                    for idx, uploaded_file in enumerate(uploaded_files):
                        if not check_rate_limit():
                            st.error("Превышен лимит: не более 10 обработок в минуту. Подождите немного.")
                            break
                        file_data = uploaded_file.getvalue()
                        with st.spinner(f"⏳ Обработка файла {uploaded_file.name} ({idx+1}/{len(uploaded_files)})..."):
                            if not check_file_size(file_data):
                                st.error(f"⚠️ Файл {uploaded_file.name} слишком большой. Максимальный размер - 1 МБ.")
                                continue
                            start_time = time.time()
                            settings = {
                                'language': 'en',
                                'optimize': optimize,
                                'use_cache': use_cache,
                                'enhance_contrast': enhance_contrast,
                                'remove_noise': remove_noise,
                                'start_time': start_time
                            }
                            result = process_image(file_data, settings)
                            processing_time = f"{time.time() - start_time:.2f} сек."
                            success = 'error' not in result
                            update_stats(success, len(file_data))
                            if success:
                                original_text = result['text']
                                detected_language = result.get('detected_language', 'auto')
                                translated_text = None
                                target_language = None
                                # Не переводим автоматически, только сохраняем результат
                                save_to_history(
                                    file_data, 
                                    original_text, 
                                    detected_language, 
                                    processing_time,
                                    translated_text,
                                    target_language
                                )
                                st.success(f"✅ {uploaded_file.name}: Текст успешно распознан за {processing_time}")
                            else:
                                st.error(f"❌ {uploaded_file.name}: {result.get('error', 'Неизвестная ошибка')}")
        
        # Предпросмотр всех загруженных файлов
        with col_preview:
            if uploaded_files:
                st.markdown("### 🖼️ Предпросмотр файлов")
                for uploaded_file in uploaded_files:
                    if uploaded_file.type.startswith('image/'):
                        st.image(uploaded_file.getvalue(), caption=uploaded_file.name, use_container_width=True, output_format="JPEG")
                    elif uploaded_file.type == 'application/pdf':
                        st.info(f"📑 PDF: {uploaded_file.name}")
                        st.markdown(f"**Размер:** {round(len(uploaded_file.getvalue()) / 1024, 2)} КБ")
        
        # Отображение результатов последней пакетной обработки
        if 'last_batch_results' in st.session_state:
            st.markdown("---")
            st.markdown("## 📝 Результаты пакетной обработки")
            for idx, res in enumerate(st.session_state['last_batch_results']):
                if not res.get('success'):
                    continue
                st.markdown(f"### 📄 {res['file_name']}")
                # Редактируемый текст
                edited_text = st.text_area(
                    f"Распознанный текст ({res['file_name']})",
                    res['original_text'],
                    height=150,
                    key=f"edit_text_{idx}",
                    label_visibility="collapsed"
                )
                # Выбор языка перевода
                lang_options = [(code, name) for code, name in TRANSLATION_LANGUAGES.items() if code != res.get('detected_language', 'en')]
                default_lang = 'en' if res.get('detected_language', 'auto') == 'ru' else 'ru'
                selected_lang = st.selectbox(
                    f"Язык перевода для {res['file_name']}",
                    options=lang_options,
                    index=[i for i, (code, _) in enumerate(lang_options) if code == default_lang][0] if any(code == default_lang for code, _ in lang_options) else 0,
                    format_func=lambda x: x[1],
                    key=f"lang_select_{idx}"
                )[0]
                # Кнопка перевода
                if st.button(f"🌐 Перевести ({res['file_name']})", key=f"translate_btn_{idx}", use_container_width=True):
                    with st.spinner("Переводим текст..."):
                        # Определяем исходный язык
                        if any(char in 'абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ' for char in edited_text):
                            source_lang = 'ru'
                        else:
                            source_lang = 'en'
                        translated_text = translate_text(edited_text, source_lang, selected_lang)
                        # Сохраняем результат в сессию для отображения
                        if 'translated_results' not in st.session_state:
                            st.session_state['translated_results'] = {}
                        st.session_state['translated_results'][idx] = {
                            'translated_text': translated_text,
                            'target_language': selected_lang,
                            'edited_text': edited_text
                        }
                        # Сохраняем в историю
                        save_to_history(
                            res['file_data'],
                            edited_text,
                            res.get('detected_language', 'auto'),
                            res['processing_time'],
                            translated_text,
                            selected_lang
                        )
                        st.success("Текст переведён и сохранён в историю!")
                # Показываем перевод, если он есть
                if 'translated_results' in st.session_state and idx in st.session_state['translated_results']:
                    t_res = st.session_state['translated_results'][idx]
                    lang_name = TRANSLATION_LANGUAGES.get(t_res['target_language'], t_res['target_language'])
                    st.text_area(
                        f"Переведённый текст ({lang_name}) — {res['file_name']}",
                        t_res['translated_text'],
                        height=150,
                        key=f"translated_text_{idx}",
                        label_visibility="collapsed"
                    )
                    # Кнопки экспорта
                    col_exp1, col_exp2, col_exp3 = st.columns(3)
                    export_content = t_res['edited_text']
                    translated_content = t_res['translated_text']
                    if col_exp1.button(f"📄 Экспорт в TXT — {res['file_name']}", key=f"export_txt_{idx}", use_container_width=True):
                        txt_file = export_to_txt(export_content + (f"\n\nПЕРЕВОД:\n{translated_content}" if translated_content else ""))
                        with open(txt_file, "rb") as file:
                            col_exp1.download_button(
                                label="⬇️ Скачать TXT",
                                data=file,
                                file_name=txt_file,
                                mime="text/plain",
                                key=f"download_txt_{idx}",
                                use_container_width=True
                            )
                    if col_exp2.button(f"📝 Экспорт в DOCX — {res['file_name']}", key=f"export_docx_{idx}", use_container_width=True):
                        docx_file = export_to_docx(export_content, translated_content)
                        with open(docx_file, "rb") as file:
                            col_exp2.download_button(
                                label="⬇️ Скачать DOCX",
                                data=file,
                                file_name=docx_file,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"download_docx_{idx}",
                                use_container_width=True
                            )
                    if col_exp3.button(f"📑 Экспорт в PDF — {res['file_name']}", key=f"export_pdf_{idx}", use_container_width=True):
                        pdf_file = export_to_pdf(export_content, translated_content)
                        with open(pdf_file, "rb") as file:
                            col_exp3.download_button(
                                label="⬇️ Скачать PDF",
                                data=file,
                                file_name=pdf_file,
                                mime="application/pdf",
                                key=f"download_pdf_{idx}",
                                use_container_width=True
                            )
                # Анализ и визуализация текста
                st.markdown("### 📊 Анализ текста")
                text_stats = analyze_text(edited_text)
                show_text_visualizations(text_stats, edited_text, key_prefix=f"{idx}_")
                st.caption(f"⏱️ Время обработки: {res['processing_time']}")
                st.divider()
    
    # Вкладка истории
    with tab2:
        st.subheader("📜 История распознавания")
        # Фильтрация и поиск
        history = load_history()
        if history:
            col_f1, col_f2, col_f3 = st.columns(3)
            search_query = col_f1.text_input("Поиск по тексту или имени файла", key="search_hist")
            lang_filter = col_f2.selectbox("Язык", options=["Все"] + list(TRANSLATION_LANGUAGES.values()), key="lang_hist")
            only_with_translation = col_f3.checkbox("Только с переводом", key="trans_hist")
            # Фильтрация
            filtered_history = []
            for item in history:
                match = True
                if search_query:
                    if search_query.lower() not in item.get('text', '').lower() and search_query.lower() not in item.get('timestamp', '').lower():
                        match = False
                if lang_filter != "Все":
                    if TRANSLATION_LANGUAGES.get(item.get('language', ''), item.get('language', '')) != lang_filter:
                        match = False
                if only_with_translation and 'translated_text' not in item:
                    match = False
                if match:
                    filtered_history.append(item)
            # Экспорт всей истории
            col_exp_hist1, col_exp_hist2, col_exp_hist3 = st.columns(3)
            if col_exp_hist1.button("⬇️ Экспорт всей истории в TXT", use_container_width=True):
                all_txt = ""
                for item in filtered_history:
                    all_txt += f"Дата: {item.get('timestamp', '')}\nЯзык: {item.get('language', '')}\nТекст:\n{item.get('text', '')}\n"
                    if 'translated_text' in item:
                        all_txt += f"Перевод:\n{item.get('translated_text', '')}\n"
                    all_txt += "\n---\n"
                txt_file = export_to_txt(all_txt)
                with open(txt_file, "rb") as file:
                    col_exp_hist1.download_button(
                        label="Скачать TXT",
                        data=file,
                        file_name=txt_file,
                        mime="text/plain",
                        use_container_width=True
                    )
            if col_exp_hist2.button("⬇️ Экспорт всей истории в DOCX", use_container_width=True):
                doc = Document()
                for item in filtered_history:
                    doc.add_heading(f"Дата: {item.get('timestamp', '')}", level=2)
                    doc.add_paragraph(f"Язык: {item.get('language', '')}")
                    doc.add_paragraph("Текст:")
                    doc.add_paragraph(item.get('text', ''))
                    if 'translated_text' in item:
                        doc.add_paragraph("Перевод:")
                        doc.add_paragraph(item.get('translated_text', ''))
                    doc.add_paragraph("---")
                docx_file = f"history_{time.strftime('%Y%m%d-%H%M%S')}.docx"
                doc.save(docx_file)
                with open(docx_file, "rb") as file:
                    col_exp_hist2.download_button(
                        label="Скачать DOCX",
                        data=file,
                        file_name=docx_file,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            if col_exp_hist3.button("⬇️ Экспорт всей истории в PDF", use_container_width=True):
                pdf = FPDF()
                pdf.set_font("Arial", size=12)
                for item in filtered_history:
                    pdf.add_page()
                    pdf.multi_cell(0, 10, f"Дата: {item.get('timestamp', '')}\nЯзык: {item.get('language', '')}\nТекст:\n{item.get('text', '')}")
                    if 'translated_text' in item:
                        pdf.ln(2)
                        pdf.multi_cell(0, 10, f"Перевод:\n{item.get('translated_text', '')}")
                    pdf.ln(2)
                    pdf.multi_cell(0, 10, "---")
                pdf_file = f"history_{time.strftime('%Y%m%d-%H%M%S')}.pdf"
                pdf.output(pdf_file)
                with open(pdf_file, "rb") as file:
                    col_exp_hist3.download_button(
                        label="Скачать PDF",
                        data=file,
                        file_name=pdf_file,
                        mime="application/pdf",
                        use_container_width=True
                    )
            # Отображение истории
            if not filtered_history:
                st.info("Нет записей по выбранным фильтрам.")
            else:
                for i, item in enumerate(filtered_history):
                    with st.expander(f"📝 **{item.get('timestamp', 'Неизвестно')}** | {item.get('language', 'Неизвестно')}"):
                        display_translation_result(item, i, show_copy_buttons=True)
        else:
            st.info("📭 История пуста. Распознайте текст, чтобы увидеть историю.")

# Функция для отображения истории перевода
def display_translation_result(item, i, show_copy_buttons=False):
    col1, col2 = st.columns([3, 1])
    with col1:
        st.markdown("#### 📄 Распознанный текст")
        st.markdown('<div class="result-container">', unsafe_allow_html=True)
        st.text_area("Оригинальный текст", item.get('text', ''), height=150, key=f"hist_text_{i}", label_visibility="collapsed")
        if show_copy_buttons:
            if st.button("📋 Копировать оригинал", key=f"copy_orig_{i}"):
                st.session_state["copy_buffer"] = item.get('text', '')
                st.toast("Текст скопирован!")
        st.markdown('</div>', unsafe_allow_html=True)
        # Отображаем переведенный текст, если он есть
        if 'translated_text' in item:
            target_lang = item.get('target_language', '')
            target_lang_name = TRANSLATION_LANGUAGES.get(target_lang, target_lang)
            st.markdown(f"#### 🌐 Перевод на {target_lang_name}")
            st.markdown('<div class="translated-container">', unsafe_allow_html=True)
            st.text_area("Переведенный текст", item.get('translated_text', ''), height=150, key=f"hist_trans_{i}", label_visibility="collapsed")
            if show_copy_buttons:
                if st.button("📋 Копировать перевод", key=f"copy_trans_{i}"):
                    st.session_state["copy_buffer"] = item.get('translated_text', '')
                    st.toast("Перевод скопирован!")
            st.markdown('</div>', unsafe_allow_html=True)
        # Визуализация текста в истории
        text_stats = analyze_text(item.get('text', ''))
        show_text_visualizations(text_stats, item.get('text', ''), key_prefix=f"hist_{i}_")
    
    with col2:
        st.caption(f"⏱️ Время обработки: {item.get('processing_time', 'Неизвестно')}")
        st.caption(f"🔤 Язык: {item.get('language', 'Неизвестно')}")
        
        # Если есть изображение, показываем его
        if 'image_path' in item:
            try:
                with open(item['image_path'], 'rb') as img_file:
                    st.image(img_file.read(), caption="Изображение", use_container_width=True, output_format="JPEG")
            except Exception:
                st.warning("Изображение недоступно")
        elif 'pdf_path' in item:
            st.info("PDF файл")
            
        # Кнопки для экспорта и анализа текста
        col_btn1, col_btn2 = st.columns(2)
        
        # Кнопка для экспорта
        if col_btn1.button("📄 Экспорт", key=f"export_btn_{i}", use_container_width=True):
            export_content = item.get('text', '')
            if 'translated_text' in item:
                export_content += f"\n\nПЕРЕВОД:\n{item.get('translated_text', '')}"
                
            export_file = export_to_txt(export_content)
            with open(export_file, "rb") as file:
                col_btn1.download_button(
                    label="⬇️ Скачать TXT",
                    data=file,
                    file_name=export_file,
                    mime="text/plain",
                    key=f"download_btn_{i}",
                    use_container_width=True
                )
        
        # Кнопка для анализа текста
        show_stats_key = f"show_stats_{i}"
        if show_stats_key not in st.session_state:
            st.session_state[show_stats_key] = False
            
        if col_btn2.button("📊 Статистика", key=f"stats_btn_{i}", use_container_width=True):
            st.session_state[show_stats_key] = not st.session_state[show_stats_key]
    
    # Отображаем статистику под колонками, если активировано
    if show_stats_key in st.session_state and st.session_state[show_stats_key]:
        text = item.get('text', '')
        st.markdown("---")
        st.markdown("### 📊 Статистика текста")
        # Получаем статистику текста
        text_stats = analyze_text(text)
        
        cols_stat = st.columns(3)
        with cols_stat[0]:
            st.markdown("##### 📝 Базовая статистика")
            st.markdown(f"""
            * **Символов:** {text_stats["chars_count"]}
            * **Слов:** {text_stats["words_count"]}
            * **Строк:** {text_stats["lines_count"]}
            * **Абзацев:** {text_stats["paragraphs_count"]}
            """)
        
        with cols_stat[1]:
            st.markdown("##### 🔤 Состав текста")
            st.markdown(f"""
            * **Букв:** {text_stats["letters_count"]}
            * **Цифр:** {text_stats["digits_count"]}
            * **Пробелов:** {text_stats["spaces_count"]}
            * **Знаков пунктуации:** {text_stats["punctuation_count"]}
            """)
        
        with cols_stat[2]:
            st.markdown("##### 📚 Частые слова")
            if text_stats["common_words"]:
                for word, count in text_stats["common_words"]:
                    st.markdown(f"* **{word}**: {count}")
            else:
                st.info("Недостаточно данных для анализа")

RATE_LIMIT = 10  # обработок в минуту
if 'rate_limit' not in st.session_state:
    st.session_state['rate_limit'] = []

def check_rate_limit():
    now = datetime.datetime.now()
    # Удаляем старые записи
    st.session_state['rate_limit'] = [t for t in st.session_state['rate_limit'] if (now - t).total_seconds() < 60]
    if len(st.session_state['rate_limit']) >= RATE_LIMIT:
        return False
    st.session_state['rate_limit'].append(now)
    return True

if __name__ == "__main__":
    main() 